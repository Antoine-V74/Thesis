"""Rewrite the speaker script docx to match the improved deck:
   - literature anchors woven into each talking point
   - full-run result numbers (+50/+100/+150 ms)
   - notes for the new slides (What I tried, NSTDB noise robustness, References)
"""
import shutil
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor

BASE = Path(r"C:\Users\antoi\OneDrive\Bureau\Master Thesis")
SRC = BASE / "Midterm_speaker_script_MNA.docx"
BAK = BASE / "Midterm_speaker_script_MNA_BACKUP.docx"
OUT = SRC

if SRC.exists() and not BAK.exists():
    shutil.copyfile(SRC, BAK)

doc = docx.Document()

# styling helpers --------------------------------------------------------
def H(text, size=15):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p

def slide(num, title):
    p = doc.add_paragraph()
    r = p.add_run(f"Slide {num} - {title}")
    r.bold = True
    r.font.size = Pt(12)
    return p

def say(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

def note(label, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    return p

# title ------------------------------------------------------------------
H("Speaker Script - Cardiac MNA Midterm", 18)
say("Main talk notes plus appendix Q&A prompts. Updated to match the improved, "
    "literature-grounded slide deck and the full uncapped benchmark numbers.")

H("Main presentation (slides 1-18)")

slide(1, "Title - Cardiac Myoneural Actuator")
say("Today I am presenting the control and safety side of the cardiac myoneural "
    "actuator project. The biological idea is a fatigue-resistant muscle actuator wrapped "
    "around the ventricle. My part asks a practical question: how do we stimulate it at the "
    "right time, and how do we avoid stimulating when the signal or rhythm is unsafe?")

slide(2, "Outline")
say("I will first explain why cardiomyoplasty is worth revisiting, then why the MNA changes "
    "the actuator problem. After that I focus on the three-layer ECG safety architecture, then "
    "the post-R stimulation-window comparison at +50, +100 and +150 milliseconds, plus a short "
    "slide on the alternatives I tried. The supplementary slides are there for deeper questions.")

slide(3, "The unmet need")
say("The clinical motivation is advanced heart failure. LVADs are life-saving, but their "
    "blood-contacting pump surface, driveline, external power, anticoagulation and infection "
    "risk are still major limitations. The ideal support helps the heart mechanically without "
    "contacting blood and without external lines.")
note("Literature", "This is exactly the rationale behind modern direct cardiac compression "
     "devices, which deliberately avoid blood contact (Roche et al. 2017, Science Translational "
     "Medicine; Bonnemain et al. 2022, Annual Review of Biomedical Engineering).")

slide(4, "Cardiomyoplasty failed as a coupled problem")
say("Dynamic cardiomyoplasty had the right broad idea: use biological tissue to compress the "
    "heart from outside. But it did not fail for one reason. Fatigue, graft degeneration, poor "
    "synchronization, diastolic interference and arrhythmia risk all interacted. I frame it as "
    "an actuator-control problem, not simply an old surgical failure.")
note("Literature", "The evidence supports this multi-factor reading: Geddes et al. (1993) "
     "showed timing mattered but was entangled with wrap tightness; Lucas et al. (1992) showed "
     "high-rate pacing preserves force but impairs relaxation; Mario et al. (1999) reported the "
     "conditioned muscle can lose more than fivefold shortening velocity and power; Salmons & "
     "Jarvis (2009) argue the field was written off too quickly.")

slide(5, "MNA changes the actuator")
say("The myoneural actuator addresses the historical weakness, fatigue. Reinnervating skeletal "
    "muscle with a sensory nerve distributes recruitment more evenly, which makes the actuator "
    "more fatigue resistant. But even a better actuator needs careful timing and safety logic: a "
    "well-timed pulse helps, a poorly timed pulse can be dangerous.")
note("Literature", "Song et al. (2025) report a myoneural actuator with about 260% greater "
     "fatigue resistance and closed-loop control demonstrated in a rodent model. That is the "
     "actuator advance my control work builds on.")

slide(6, "Controller scope")
say("For this midterm the objective is not chronic animal support. It is to build the safety and "
    "timing framework needed before that experiment. The controller must detect R peaks, decide "
    "whether the beat is safe, and then permit or inhibit stimulation. The comparison later tests "
    "how much post-R signal that decision needs.")
note("Literature", "This maps onto a hierarchical controller: a fast hardware-friendly trigger "
     "loop, a beat-to-beat decision loop, and a slow supervisory loop - the structure in the "
     "project's ECG-gating notes, and consistent with adaptive assist-control work (Magkoutas et "
     "al. 2022).")

slide(7, "Layered safety architecture")
say("The architecture is intentionally layered. Layer 1 is fast and deterministic. Layer 2 is a "
    "personalized feature gate calibrated on healthy ECG. Layer 3 is a research anomaly layer. The "
    "rule is conservative: stimulation is permitted only when the required layers agree. Any "
    "uncertainty, missing data or runtime issue defaults to inhibit. Machine learning can only "
    "veto, it can never command stimulation by itself.")

slide(8, "Layer 1 threshold algorithm")
say("Layer 1 is a causal threshold detector. It filters the ECG, builds a local energy and slope "
    "envelope, applies an adaptive threshold, confirms the local peak after a short descent, then "
    "passes candidates to the RR supervisor. It does not diagnose arrhythmia; it provides a fast, "
    "explainable trigger stream.")

slide(9, "RR supervisor")
say("The RR supervisor turns detections into safe candidates. During calibration no stimulation "
    "is allowed. In running mode accepted beats update the RR reference. If RR intervals become "
    "implausible, the system enters recovery or inhibition. This matters because duplicate "
    "detections and missed beats are exactly the errors that could schedule stimulation at the "
    "wrong time.")

slide(10, "Layer 2 safety gate")
say("Layer 2 asks a different question: should this beat be stimulated? It uses timing, "
    "morphology, signal-quality, energy and wavelet features. The gate is calibrated on the "
    "record's own healthy baseline, then inhibits beats that deviate too far. Concretely it is a "
    "personalized anomaly gate: robust z-scores plus a Ledoit-Wolf Mahalanobis distance against "
    "the healthy baseline. It is not a universal arrhythmia classifier; it is a personalized "
    "safety veto.")

slide(11, "Layer 3")
say("Layer 3 is deliberately a research layer. It uses self-supervised representation learning "
    "and anomaly scoring, but it is not a hidden controller. Until it is prospectively validated "
    "it can only add conservative evidence, and failure or uncertainty means inhibit.")

slide(12, "Layer 1 result")
say("The fast causal detector is strong enough to support the rest of the work. On clean normal "
    "sinus ECG it reaches very high sensitivity and PPV; in the cleaner NSTDB subset (SNR at or "
    "above +12 dB) it keeps 100% sensitivity with only a few milliseconds of confirmation delay. "
    "The remaining issues are noise, duplicates and recovery behavior, which is exactly what the "
    "supervisor and Layer 2 are for.")

slide(13, "Post-R window comparison (+50/+100/+150 ms)")
say("This is the key result slide, and these are full uncapped cross-dataset runs. At +50 ms "
    "healthy permit collapses to about 30 percent - the morphology is not complete yet - even "
    "though abnormal inhibition stays near 96 percent and false permit stays low. At +100 ms "
    "healthy permit jumps to about 69 percent, abnormal inhibition is about 97 percent, and false "
    "permit is about 3 percent. +150 ms is marginally higher on permit but false permit creeps up "
    "and latency grows.")
note("Literature", "This fits the physiology: Geddes et al. (1993) found a subject-specific "
     "optimal R-delay around 58 ms (range 40-80 ms), and the cardiac electromechanical window is "
     "only about 60-100 ms, so +100 ms being the first robust point is consistent.")

slide(14, "Interpretation")
say("The result changes what I emphasize. I do not present the centered plus-minus 2.5 second "
    "result, because that is non-causal and only an offline upper bound. The real deployment "
    "question is how much future ECG after the R peak the decision needs. The current answer is "
    "+100 ms.")

slide(15, "What I tried (before choosing +100 ms)")
say("Before settling on +100 ms I compared four design axes. First, windowing: I started from a "
    "centered plus-minus 2.5 second oracle, which is non-causal and only an upper bound, then "
    "moved to a causal window. Second, decision timing: gated, next-beat permit, and a stateful "
    "risk carry-over; gated is the cleanest, the others trade latency for an earlier permit. "
    "Third, the feature set: all features versus signal-only - all features keep abnormal "
    "inhibition high, signal-only mainly rejects noise. Fourth, thresholds: I swept the "
    "Mahalanobis scale and the RR coupling threshold to push false permit down. The defensible "
    "operating point is +100 ms, causal, gated, all features, with a tuned threshold.")

slide(16, "Next steps")
say("The next phase connects the ECG safety decision to mechanical benefit: bring pressure data "
    "into the TDT pipeline, calibrate per animal, measure total stimulation delay, and compare "
    "fixed versus adaptive stimulation timing using pressure or flow response.")
note("Literature", "Pressure-based labelling is what the historical work actually used to define "
     "the useful window (Geddes et al. 1993). I should expect a noisy mechanical target - "
     "automated pre-ejection-period extraction has a real error budget (Richer et al. 2025) - and "
     "intracardiac impedance is a candidate chronic surrogate (Osswald et al. 2000).")

slide(17, "Take-home message")
say("MNA makes biological direct compression worth revisiting, but the controller is what makes "
    "it testable. I built a safety-first three-layer ECG architecture, showed the threshold/RR "
    "layer supports causal validation, and found that +100 ms is the first robust post-R decision "
    "window. Next is connecting that safe trigger to measured mechanical benefit.")

slide(18, "Questions")
say("I would stop here for the main talk. For detail I can use the appendix slides on the "
    "literature, threshold defaults, RR supervisor, Layer 2 feature groups, full numeric table, "
    "noise robustness, operating curves, limitations and references.")

# supplementary ----------------------------------------------------------
H("Supplementary slides (19-30) - use on demand")
note("Slide 27, NSTDB noise robustness",
     "At SNR at or above +12 dB the gate inhibits 100 percent of abnormal beats with zero false "
     "permit. It only degrades below 0 dB, where the system should lean on Layer-1 signal-quality "
     "inhibition - uncertainty defaults to inhibit.")
note("Slide 28, operating curves",
     "Use if asked how the chosen threshold trades healthy availability against abnormal "
     "inhibition; the Mahalanobis and coupling-threshold sweeps show how to drive false permit "
     "down.")
note("Slide 29, limitations",
     "Human-to-rat transfer needs recalibration; part of the deployment gap is Layer-1 trigger "
     "quality, not only Layer 2; narrow SVT can look too normal. False permit is the critical "
     "failure mode, so uncertainty defaults to inhibit.")
note("Slide 30, references", "Full reference list, grounding every design choice in a source.")

# appendix Q&A -----------------------------------------------------------
H("Appendix Q&A prompts")

def qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q)
    r.bold = True
    doc.add_paragraph("A: " + a)

qa("Why not just use the centered 2.5 s windows?",
   "They use future samples and are an offline upper bound - useful for debugging, not for "
   "deployment claims.")
qa("Why is +50 ms so poor on healthy permit?",
   "The QRS tail and early ST morphology are incomplete, so template and morphology features are "
   "unreliable and the gate becomes conservative. Importantly, abnormal inhibition and false "
   "permit stay good even at +50 ms - it is availability that suffers.")
qa("Why is +100 ms acceptable if stimulation timing may need to be earlier?",
   "Layer 2 can run as a next-beat permit state: beat N computes the safety state used for "
   "stimulation on beat N+1, while immediate hard vetoes can still fire on the same beat. "
   "Physiologically the optimal R-delay is also in the tens-of-milliseconds range (Geddes et al. "
   "1993), so a one-beat-latent permit is compatible with the mechanical window.")
qa("What is the most important safety metric?",
   "False permit, because permitting stimulation on an unsafe beat could be pro-arrhythmic. False "
   "inhibit reduces therapy availability but is safer.")
qa("How do I know these numbers are trustworthy?",
   "They come from full uncapped cross-dataset runs (MIT-BIH, NSTDB, SVDB) in the deployment mode "
   "(fast_causal_gated, all features), not a quick capped sweep, and the NSTDB results are broken "
   "down by SNR.")
qa("What still needs validation?",
   "Rat ECG calibration, noise/SQI thresholds, total hardware and muscle-activation delay, and "
   "pressure-based confirmation of mechanical benefit.")

try:
    doc.save(str(OUT))
    print("Saved script:", OUT)
except PermissionError:
    OUT = BASE / "Midterm_speaker_script_MNA_v2.docx"
    doc.save(str(OUT))
    print("Original locked. Saved script to:", OUT)
